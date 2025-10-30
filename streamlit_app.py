from __future__ import annotations

import copy
import json
import os
import unicodedata
import tempfile
from dataclasses import dataclass
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

import streamlit as st
from docx import Document
from openai import OpenAI


@dataclass
class MeetingAnalysis:
    language: str
    transcript: str
    summary_points: List[str]
    keywords: List[str]
    action_items: List[str]
    conversation: List[Dict[str, Any]]
    coverage_ratio: float
    similarity: float
    conversation_fallback: bool


MEETING_ANALYSIS_JSON_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": ["language", "summary", "conversation"],
    "properties": {
        "language": {"type": "string"},
        "summary": {
            "type": "object",
            "additionalProperties": False,
            "required": ["key_points", "keywords", "action_items"],
            "properties": {
                "key_points": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "keywords": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "action_items": {
                    "type": "array",
                    "items": {"type": "string"},
                },
            },
        },
        "conversation": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["speaker", "original_text", "traditional_chinese", "notes"],
                "properties": {
                    "speaker": {"type": "string"},
                    "original_text": {"type": "string"},
                    "traditional_chinese": {"type": "string"},
                    "notes": {"type": ["string", "null"]},
                },
            },
        },
    },
}


def _load_openai_client(explicit_key: Optional[str] = None) -> OpenAI:
    api_key: Optional[str] = explicit_key or os.getenv("OPENAI_API_KEY")
    try:
        if "OPENAI_API_KEY" in st.secrets:
            api_key = st.secrets["OPENAI_API_KEY"]
    except (RuntimeError, FileNotFoundError):
        pass

    if not api_key:
        st.error("請在側邊欄輸入 OpenAI API Key，或設定環境變數 OPENAI_API_KEY。")
        st.stop()
    return OpenAI(api_key=api_key)


def _transcribe_audio(client: OpenAI, file_path: Path, model: str, language_hint: Optional[str]) -> Dict[str, Any]:
    if model == "whisper-1":
        response_format = "verbose_json"
    else:
        response_format = "json"

    kwargs: Dict[str, Any] = {"language": language_hint} if language_hint else {}
    with file_path.open("rb") as audio_file:
        result = client.audio.transcriptions.create(
            model=model,
            file=audio_file,
            response_format=response_format,
            **kwargs,
        )

    if hasattr(result, "model_dump"):
        return result.model_dump()
    return dict(result)


def _normalize_text_for_comparison(text: str) -> str:
    return "".join(
        ch
        for ch in text
        if not unicodedata.category(ch).startswith(("P", "Z", "C"))
    )


def _conversation_metrics(conversation: List[Dict[str, Any]], transcript: str) -> Dict[str, float]:
    transcript_norm = _normalize_text_for_comparison(transcript)
    conversation_norm = _normalize_text_for_comparison(
        "".join(turn.get("original_text", "") for turn in conversation)
    )
    if not transcript_norm:
        return {"coverage_ratio": 1.0, "similarity": 1.0}

    coverage_ratio = min(len(conversation_norm) / len(transcript_norm), 1.0)
    similarity = SequenceMatcher(None, transcript_norm, conversation_norm).ratio() if conversation_norm else 0.0
    return {"coverage_ratio": coverage_ratio, "similarity": similarity}


def _extract_structured_payload(response: Any) -> Dict[str, Any]:
    payload: Dict[str, Any] = {}
    text_chunks: List[str] = []

    for item in getattr(response, "output", []) or []:
        content_items = getattr(item, "content", [])
        if not isinstance(content_items, list):
            continue
        for piece in content_items:
            piece_type = ""
            piece_text = ""
            piece_json: Optional[Dict[str, Any]] = None
            if isinstance(piece, dict):
                piece_type = piece.get("type", "")
                piece_text = piece.get("text", "")
                if "json" in piece:
                    piece_json = piece.get("json")
            else:
                piece_type = getattr(piece, "type", "")
                piece_text = getattr(piece, "text", "")
                piece_json = getattr(piece, "json", None)

            if not payload and piece_type == "output_json" and isinstance(piece_json, dict):
                payload = piece_json
            elif piece_type == "output_text" and piece_text:
                text_chunks.append(piece_text)

    if not payload:
        text_output = ""
        if hasattr(response, "output_text"):
            text_output = getattr(response, "output_text") or ""
        if not text_output and text_chunks:
            text_output = "".join(text_chunks).strip()
        if text_output:
            try:
                payload = json.loads(text_output)
            except json.JSONDecodeError:
                payload = {}

    return payload


def _sanitize_conversation(
    conversation: List[Dict[str, Any]],
    target_language: str,
) -> List[Dict[str, Any]]:
    sanitized: List[Dict[str, Any]] = []
    for index, turn in enumerate(conversation or []):
        speaker = turn.get("speaker") or f"講者{chr(65 + index)}"
        original = str(turn.get("original_text") or "").strip()
        translated = str(turn.get("traditional_chinese") or "").strip()
        notes = turn.get("notes")
        if notes is None or notes == "":
            notes_value: Optional[str] = None
        else:
            notes_value = str(notes).strip()

        if not translated:
            translated = original if target_language == "繁體中文" else original

        sanitized.append(
            {
                "speaker": str(speaker),
                "original_text": original,
                "traditional_chinese": translated,
                "notes": notes_value,
            }
        )

    return sanitized


def _fallback_conversation(
    transcript_text: str,
    target_language: str,
) -> List[Dict[str, Any]]:
    original = transcript_text.strip()
    return [
        {
            "speaker": "講者A",
            "original_text": original,
            "traditional_chinese": original if target_language == "繁體中文" else original,
            "notes": None,
        }
    ]


def _ensure_str_list(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(item) for item in value if item not in (None, "")]
    if value in (None, ""):
        return []
    return [str(value)]


def _analyse_meeting(
    client: OpenAI,
    transcript_text: str,
    target_language: str,
    model: str,
) -> MeetingAnalysis:
    base_messages = [
        {
            "role": "system",
            "content": [
                {
                    "type": "input_text",
                    "text": (
                        "You are an assistant that rewrites meeting notes. "
                        "Break transcripts into discrete speaker turns and infer different speakers "
                        "when the language indicates a change in perspective (e.g. questions vs answers). "
                        "Return structured Traditional Chinese output describing the conversation. "
                        "Keep faithful to the original speech, do not abridge or skip any part of the transcript, "
                        "and include concise action items."
                    ),
                }
            ],
        },
        {
            "role": "user",
            "content": [
                {
                    "type": "input_text",
                    "text": (
                        "Transcript:\n"
                        f"{transcript_text}\n\n"
                        "請將重點摘要、關鍵字、待辦事項整理成繁體中文，並生成 Speaker A/B/... 的對話紀錄。"
                        f"翻譯語言請使用 {target_language}，必要時合併破碎語句，但不得捏造內容。"
                        "輸出請採 JSON 格式，欄位包含 language、summary.key_points、summary.keywords、summary.action_items、conversation。"
                        "conversation 每項需含 speaker、original_text、traditional_chinese（翻譯）、notes（可選）。"
                        "依照語意切分段落，conversation 的 original_text 必須覆蓋完整逐字稿內容（僅能做斷句與微幅標點修正，不得刪減）。"
                        "若看出多位講者，請適當切分段落並標註不同講者；無法確定時可使用『講者A/B/...』。"
                        "notes 欄位可為 null；若內容清楚可省略補充說明。"
                    ),
                }
            ],
        },
    ]

    payload: Dict[str, Any] = {}
    conversation: List[Dict[str, Any]] = []
    metrics = {"coverage_ratio": 0.0, "similarity": 0.0}
    conversation_fallback = False

    max_attempts = 3
    for attempt in range(max_attempts):
        messages = copy.deepcopy(base_messages)
        if attempt > 0:
            coverage_note = (
                "上一次輸出的 conversation 未完整覆蓋逐字稿內容。"
                f"覆蓋率約為 {metrics['coverage_ratio']:.0%}，相似度約為 {metrics['similarity']:.0%}。"
                "請重新切分逐字稿，逐段輸出原文與翻譯，確保 original_text 串接後"
                "能與逐字稿內容等長且語意相符，不得遺漏任何部分：\n"
                f"{transcript_text}"
            )
            messages.append(
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": coverage_note,
                        }
                    ],
                }
            )

        response = client.responses.create(
            model=model,
            text={
                "format": {
                    "type": "json_schema",
                    "name": "meeting_analysis",
                    "schema": MEETING_ANALYSIS_JSON_SCHEMA,
                }
            },
            input=messages,
        )

        payload = _extract_structured_payload(response) or {}
        if "summary" not in payload or not isinstance(payload["summary"], dict):
            payload["summary"] = {}

        conversation = _sanitize_conversation(payload.get("conversation", []), target_language)
        payload["conversation"] = conversation
        metrics = _conversation_metrics(conversation, transcript_text)

        if conversation and metrics["coverage_ratio"] >= 0.95 and metrics["similarity"] >= 0.85:
            break
    else:
        conversation = _fallback_conversation(transcript_text, target_language)
        payload.setdefault("summary", {})
        payload["conversation"] = conversation
        payload.setdefault("language", "unknown")
        conversation_fallback = True
        metrics = _conversation_metrics(conversation, transcript_text)

    summary = payload.get("summary", {})

    return MeetingAnalysis(
        language=payload.get("language", "unknown"),
        transcript=transcript_text,
        summary_points=_ensure_str_list(summary.get("key_points")),
        keywords=_ensure_str_list(summary.get("keywords")),
        action_items=_ensure_str_list(summary.get("action_items")),
        conversation=conversation,
        coverage_ratio=metrics["coverage_ratio"],
        similarity=metrics["similarity"],
        conversation_fallback=conversation_fallback,
    )


def _render_summary(analysis: MeetingAnalysis) -> None:
    st.subheader("摘要重點")
    if analysis.summary_points:
        for point in analysis.summary_points:
            st.markdown(f"- {point}")
    else:
        st.info("無摘要內容。")

    if analysis.action_items:
        st.subheader("待辦事項")
        for item in analysis.action_items:
            st.markdown(f"- ✅ {item}")

    if analysis.keywords:
        st.subheader("關鍵字")
        st.write("、".join(analysis.keywords))


def _render_conversation(analysis: MeetingAnalysis) -> None:
    st.subheader("對話紀錄")
    if not analysis.conversation:
        st.info("無對話資料。")
        return

    if analysis.conversation_fallback:
        st.warning(
            "僅能產生單一段落的對話資料；原文已完整保留，翻譯欄位與原文相同，建議手動調整。"
        )
    elif analysis.coverage_ratio < 0.95:
        st.warning(
            f"對話覆蓋率約為 {analysis.coverage_ratio:.0%}，可能仍有缺漏，請人工複核。"
        )

    for turn in analysis.conversation:
        speaker = turn.get("speaker", "Speaker")
        original = turn.get("original_text", "")
        translated = turn.get("traditional_chinese", "")
        notes = turn.get("notes")

        st.markdown(f"**{speaker}**")
        if original:
            st.markdown(f"- 原文：{original}")
        if translated and translated != original:
            st.markdown(f"- 翻譯：{translated}")
        if notes:
            st.markdown(f"- 備註：{notes}")
        st.divider()


def _render_downloads(analysis: MeetingAnalysis, raw_transcript: Dict[str, Any]) -> None:
    export_payload = {
        "language": analysis.language,
        "summary_points": analysis.summary_points,
        "keywords": analysis.keywords,
        "action_items": analysis.action_items,
        "conversation": analysis.conversation,
        "metrics": {
            "conversation_coverage_ratio": analysis.coverage_ratio,
            "conversation_similarity": analysis.similarity,
            "conversation_fallback": analysis.conversation_fallback,
        },
        "raw_transcription": raw_transcript,
    }
    st.download_button(
        label="下載 JSON",
        data=json.dumps(export_payload, ensure_ascii=False, indent=2),
        mime="application/json",
        file_name="meeting_analysis.json",
    )

    document = Document()
    document.add_heading("會議記錄分析器・普賢宮", level=1)
    document.add_paragraph(f"語言偵測：{analysis.language}")
    if analysis.conversation_fallback:
        document.add_paragraph("注意：僅能產生單一段落的對話資料，請人工檢查並補充。")
    elif analysis.coverage_ratio < 0.95:
        document.add_paragraph(
            f"注意：系統估計對話覆蓋率約為 {analysis.coverage_ratio:.0%}，建議人工複核。"
        )

    if analysis.summary_points:
        document.add_heading("摘要", level=2)
        for point in analysis.summary_points:
            document.add_paragraph(point, style="List Bullet")

    if analysis.action_items:
        document.add_heading("待辦事項", level=2)
        for item in analysis.action_items:
            document.add_paragraph(item, style="List Number")

    if analysis.keywords:
        document.add_heading("關鍵字", level=2)
        document.add_paragraph("、".join(analysis.keywords))

    document.add_heading("對話紀錄", level=2)
    for turn in analysis.conversation:
        speaker = turn.get("speaker", "Speaker")
        original = turn.get("original_text", "")
        translated = turn.get("traditional_chinese", "")
        notes = turn.get("notes")

        para = document.add_paragraph()
        para.add_run(f"{speaker}\n").bold = True
        if original:
            para.add_run(f"原文：{original}\n")
        if translated and translated != original:
            para.add_run(f"翻譯：{translated}\n")
        if notes:
            para.add_run(f"備註：{notes}\n")

    document.add_heading("完整逐字稿", level=2)
    document.add_paragraph(analysis.transcript)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="下載 Word (.docx)",
        data=buffer,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_name="meeting_analysis.docx",
    )


def main() -> None:
    st.set_page_config("會議記錄分析器・普賢宮", layout="wide")
    st.title("會議記錄分析器・普賢宮")
    st.caption("上傳台語或國語錄音，透過 OpenAI API 進行轉錄、翻譯與摘要。")

    with st.sidebar:
        st.header("設定")
        default_key = st.session_state.get("openai_api_key", "")
        sidebar_key = st.text_input(
            "OpenAI API Key",
            value=default_key,
            type="password",
            help="輸入後僅儲存在瀏覽器 Session，不會寫入硬碟。",
        )
        if sidebar_key:
            st.session_state["openai_api_key"] = sidebar_key
            os.environ["OPENAI_API_KEY"] = sidebar_key

        transcription_model = st.selectbox(
            "轉錄模型",
            options=[
                "gpt-4o-mini-transcribe",
                "gpt-4o-transcribe",
                "whisper-1",
            ],
            index=0,
        )
        analysis_model = st.selectbox(
            "分析模型",
            options=[
                "gpt-4o-mini",
                "gpt-4o",
                "o4-mini",
            ],
            index=0,
        )
        target_language = st.selectbox(
            "翻譯輸出語言",
            options=["繁體中文", "English", "日本語"],
            index=0,
        )
        language_hint_option = st.selectbox(
            "音訊主要語言",
            options=["自動偵測", "國語 (zh)", "台語 (nan)"],
            index=0,
        )
        language_hint = None
        if language_hint_option == "國語 (zh)":
            language_hint = "zh"
        elif language_hint_option == "台語 (nan)":
            language_hint = "nan"

    client = _load_openai_client(st.session_state.get("openai_api_key"))

    uploaded_file = st.file_uploader(
        "上傳音檔", type=["mp3", "wav", "m4a", "aac", "flac", "ogg"]
    )

    if not uploaded_file:
        st.info("請先選擇音檔。")
        return

    if st.button("開始分析", type="primary"):
        with st.spinner("轉錄與分析中，請稍候..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp:
                tmp.write(uploaded_file.read())
                temp_path = Path(tmp.name)

            try:
                transcription = _transcribe_audio(
                    client, temp_path, transcription_model, language_hint
                )
                transcript_text = transcription.get("text", "")
                if not transcript_text:
                    st.error("未取得轉錄結果，請稍後再試。")
                    return

                analysis = _analyse_meeting(
                    client=client,
                    transcript_text=transcript_text,
                    target_language=target_language,
                    model=analysis_model,
                )
            except Exception as exc:
                st.error(f"處理過程發生錯誤：{exc}")
                return
            finally:
                temp_path.unlink(missing_ok=True)

        st.success("分析完成！")
        _render_summary(analysis)
        _render_conversation(analysis)
        if analysis.conversation:
            st.caption(
                f"對話覆蓋率：約 {analysis.coverage_ratio:.0%}｜相似度：約 {analysis.similarity:.0%}"
            )

        with st.expander("查看原始逐字稿"):
            st.write(analysis.transcript)

        copy_blocks: List[str] = []
        for turn in analysis.conversation:
            speaker = turn.get("speaker", "Speaker")
            translated = turn.get("traditional_chinese") or turn.get("original_text", "")
            copy_blocks.append(f"{speaker}: {translated}")

        st.subheader("快速複製")
        st.text_area(
            "可直接複製以下內容：",
            value="\n".join(copy_blocks) if copy_blocks else analysis.transcript,
            height=200,
        )

        _render_downloads(analysis, transcription)


if __name__ == "__main__":
    main()
